"""
Word 报告生成器 —— 模板替换与合并。
"""

import os
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt


class WordReportGenerator:
    """处理 Word 模板替换与合并 (支持追加模式)"""

    @staticmethod
    def generate_report(template_path, output_path, data_dict):
        """生成单个临时报告文件"""
        try:
            doc = Document(template_path)
            for table in doc.tables:
                table.autofit = False
                table.allow_autofit = False

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        WordReportGenerator._replace_text_in_cell(cell, data_dict)
            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)

    @staticmethod
    def _replace_text_in_cell(cell, data_dict):
        if not cell.text:
            return
        for key, value in data_dict.items():
            if key in cell.text:
                cell.text = cell.text.replace(key, value)
                if "CODE" in key or "代码" in key:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Consolas"
                            run.font.size = Pt(9)

    @staticmethod
    def merge_reports(temp_file_list, final_output_path):
        """
        合并多个临时报告到最终文件。
        如果 final_output_path 已存在则追加，否则新建。
        """
        if not temp_file_list:
            return False, "没有文件可合并"

        try:
            if os.path.exists(final_output_path):
                master_doc = Document(final_output_path)
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                master_doc.add_paragraph("\n")
                p = master_doc.add_paragraph(f"=== 新增批次: {timestamp} ===")
                p.alignment = 1
                master_doc.add_paragraph("\n")
                files_to_process = temp_file_list
            else:
                master_doc = Document(temp_file_list[0])
                files_to_process = temp_file_list[1:]

            for sub_file_path in files_to_process:
                sub_doc = Document(sub_file_path)
                master_doc.add_paragraph("\n")
                for element in sub_doc.element.body:
                    master_doc.element.body.append(element)

            master_doc.save(final_output_path)
            return True, "合并成功"
        except Exception as e:
            return False, str(e)
