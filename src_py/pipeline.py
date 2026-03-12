from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QTableWidget, QTableWidgetItem, QApplication,
    QHeaderView, QProgressBar, QHBoxLayout, QLabel, QLineEdit, QFileDialog, QFormLayout,
    QPushButton, QTextEdit, QMessageBox, QGroupBox, QListWidget, QListWidgetItem, QCheckBox
)
from PyQt5.QtCore import Qt, QTimer
from net_workers import GenerationThread
import os
import re
import shutil
from urllib.parse import unquote
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime

# 加载含有参考代码的excel文件
class RefCodeManager:
    def __init__(self, excel_path):
        self.ref_data = {} # Key: ID, Value: Code
        if excel_path and os.path.exists(excel_path):
            self.load_excel(excel_path)

    def load_excel(self, path):
        try:
            df = pd.read_excel(path, engine='openpyxl')
            # 自动规整列名
            df.columns = [str(col).strip() for col in df.columns]
            
            # 模糊匹配列名
            id_col = next((c for c in df.columns if "ID" in c or "id" in c), None)
            code_col = next((c for c in df.columns if "代码" in c or "Code" in c), None)

            if id_col and code_col:
                for _, row in df.iterrows():
                    req_id = str(row[id_col]).strip()
                    code = str(row[code_col])
                    if code.lower() == 'nan': code = "无参考代码"
                    self.ref_data[req_id] = code
                print(f"参考代码库加载成功: {len(self.ref_data)} 条")
        except Exception as e:
            print(f"参考代码加载失败: {e}")

    def get_code(self, req_id):
        return self.ref_data.get(req_id, "无 (未在参考库中找到对应ID)")

class WordReportGenerator:
    """处理 Word 模板替换与合并 (支持追加模式)"""   
    @staticmethod
    def generate_report(template_path, output_path, data_dict):
        """生成单个临时报告文件"""
        try:
            doc = Document(template_path)
            for table in doc.tables:
                table.autofit = False
                table.allow_autofit = False

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        WordReportGenerator.replace_text_in_cell(cell, data_dict)
            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)

    @staticmethod
    def replace_text_in_cell(cell, data_dict):
        if not cell.text: return
        for key, value in data_dict.items():
            if key in cell.text:
                cell.text = cell.text.replace(key, value)
                # 可选：简单的字体修正
                if "CODE" in key:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)

    @staticmethod
    def merge_reports(temp_file_list, final_output_path):
        """
        核心逻辑：
        1. 如果 final_output_path 存在 -> 追加模式
        2. 如果 final_output_path 不存在 -> 新建模式
        """
        if not temp_file_list:
            return False, "没有文件可合并"

        try:
            # 判断是否追加
            if os.path.exists(final_output_path):
                # === 追加模式 ===
                master_doc = Document(final_output_path)
                
                # 插入批次分隔符
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                master_doc.add_paragraph("\n")
                p = master_doc.add_paragraph(f"=== 新增批次: {timestamp} ===")
                p.alignment = 1 # 居中
                master_doc.add_paragraph("\n")
                
                files_to_process = temp_file_list
            else:
                # === 新建模式 ===
                # 用第一个临时文件做底板
                master_doc = Document(temp_file_list[0])
                files_to_process = temp_file_list[1:]

            # 循环追加剩余文件
            for sub_file_path in files_to_process:
                sub_doc = Document(sub_file_path)
                
                # 两个表格之间插入回车
                master_doc.add_paragraph("\n") 
                
                # 搬运内容
                for element in sub_doc.element.body:
                    master_doc.element.body.append(element)

            master_doc.save(final_output_path)
            return True, "合并成功"
        except Exception as e:
            return False, str(e)

class FileDragTextEdit(QTextEdit):
    """自定义文本框：拖入多个文件时，自动换行显示路径，而不是连在一起"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)

    def insertFromMimeData(self, source):
        if source.hasUrls():
            paths = []
            for url in source.urls():
                if url.isLocalFile():
                    paths.append(url.toLocalFile())
            
            if paths:
                self.insertPlainText("\n".join(paths) + "\n")
        else:
            super().insertFromMimeData(source)


class VariableManager:
    """管理变量Excel表，提供检索功能"""
    def __init__(self, excel_path):
        self.df = pd.DataFrame()
        self.is_loaded = False
        self.load_excel(excel_path)

    def load_excel(self, excel_path):
        if not os.path.exists(excel_path):
            return
        try:
            # 读取 Excel
            self.df = pd.read_excel(excel_path, engine='openpyxl')
            # 填充空值
            self.df.fillna("", inplace=True)
            # 标准化列名映射 (防止用户Excel表头有微小差异)
            self.df['search_index'] = self.df.apply(
                lambda x: f"{x.get('信号名称','')} {x.get('信号ID（变量名）','')} {x.get('值定义','')}", axis=1
            )
            self.is_loaded = True
            print(f"变量表加载成功，共 {len(self.df)} 条数据")
        except Exception as e:
            print(f"变量表加载失败: {e}")

    def search_relevant_vars(self, requirement_text, top_k=10):
        if not self.is_loaded or self.df.empty:
            return []

        # 简单的切词 (按非字字符切割)
        req_keywords = set(re.split(r'[^\w]+', requirement_text))
        req_keywords = {k for k in req_keywords if len(k) > 1} 

        results = []
        for index, row in self.df.iterrows():
            score = 0
            search_content = str(row['search_index'])
            
            for kw in req_keywords:
                if kw in search_content:
                    score += 1
            
            if score > 0:
                var_info = (
                    f"- 名称: {row.get('信号名称','')}, "
                    f"ID: {row.get('信号ID（变量名）','')}, "
                    f"类型: {row.get('数据类型','')}, "
                    f"定义: {row.get('值定义','')}"
                )
                results.append((score, var_info))

        results.sort(key=lambda x: x[0], reverse=True)
        return [r[1] for r in results[:top_k]]


class BatchTask:
    """定义单个流水线任务"""
    def __init__(self, raw_text, parsed_data):
        self.raw_text = raw_text
        self.id = parsed_data.get("req_id", "Unknown")
        self.name = parsed_data.get("req_name", "Unknown")
        self.output_rel_path = parsed_data.get("output_file", "").strip()
        self.content = parsed_data.get("req_content", "")
        self.vars = parsed_data.get("req_vars", "")
        self.ref_code = parsed_data.get("req_ref_code", "") # 新增参考代码字段
        
        #ad-hoc: 如果parsed_data里有 target_files 则直接用，否则置空
        self.target_files = parsed_data.get("target_files", [])
        
        self.status = "等待中" 
        self.result_code = ""
        self.error_msg = ""

        self.generated_clean_code = ""
        self.used_global_vars = ""

class PipelineTab(QWidget):
    """流水线处理 Tab"""
    def __init__(self, app_parent):
        super().__init__()
        self.app = app_parent 
        self.tasks = []
        self.current_task_index = -1
        self.is_running = False
        
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        input_group = QGroupBox("1. 批量需求导入")
        ig_layout = QVBoxLayout()
        
        hint = QLabel("💡 提示：支持拖入 .txt/.md (正则解析) 或 .xlsx (表头: 需求ID, 需求名称, 输出文件, 需求内容, 参考代码)。")
        hint.setStyleSheet("color: #666; font-size: 12px;")
        ig_layout.addWidget(hint)

        # ===使用自定义控件 FileDragTextEdit ===
        self.batch_input_edit = FileDragTextEdit() 
        self.batch_input_edit.setPlaceholderText("在此粘贴文本，或直接拖入 .xlsx / .txt 文件...")
        ig_layout.addWidget(self.batch_input_edit)
        # ============================================
        
        btn_layout = QHBoxLayout()
        self.parse_btn = QPushButton("🔍 智能解析")
        self.parse_btn.clicked.connect(self.parse_tasks)
        btn_layout.addWidget(self.parse_btn)
        
        self.clear_btn = QPushButton("清空列表")
        self.clear_btn.clicked.connect(self.clear_tasks)
        btn_layout.addWidget(self.clear_btn)
        
        ig_layout.addLayout(btn_layout)
        input_group.setLayout(ig_layout)
        layout.addWidget(input_group, 3)

        #=== RAG ===
        rag_group = QGroupBox("2. RAG 增强配置")
        rag_layout = QHBoxLayout()
        
        self.pipe_use_rag_cb = QCheckBox("启用流水线 RAG")
        self.pipe_use_rag_cb.setChecked(True) # 默认开启
        rag_layout.addWidget(self.pipe_use_rag_cb)
        
        # 知识库列表 
        self.pipe_kb_list = QListWidget()
        self.pipe_kb_list.setMaximumHeight(50)
        self.pipe_kb_list.setStyleSheet("border: 1px solid #ccc; border-radius: 4px;")
        
        # 加载现有知识库
        all_kbs = self.app.rag_manager.knowledge_bases
        for kb in all_kbs:
            item = QListWidgetItem(kb)
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            if "GJB" in kb:
                item.setCheckState(Qt.Checked)
            else:
                item.setCheckState(Qt.Unchecked)
            self.pipe_kb_list.addItem(item)
            
        rag_layout.addWidget(self.pipe_kb_list)
        rag_group.setLayout(rag_layout)
        layout.addWidget(rag_group, 1)

        var_group = QGroupBox("2.1 变量知识库 (Excel)")
        var_layout = QHBoxLayout()
        
        self.var_path_edit = QLineEdit(self.app.config.variable_excel_path)
        self.var_path_edit.setPlaceholderText("选择包含 '信号名称/ID/类型' 的 Excel 文件...")
        
        var_btn = QPushButton("📂 选择变量表")
        var_btn.clicked.connect(self.select_variable_file)
        
        var_layout.addWidget(QLabel("变量表:"))
        var_layout.addWidget(self.var_path_edit)
        var_layout.addWidget(var_btn)
        
        var_group.setLayout(var_layout)
        layout.addWidget(var_group, 0)

        # === 新增：报告生成配置 ===
        report_group = QGroupBox("2.2 测试报告配置 (Word)")
        report_layout = QFormLayout()
        
        # 模板选择
        self.word_template_path = QLineEdit()
        self.word_template_path.setPlaceholderText("选择包含表格和占位符的 .docx 模板...")
        browse_tpl_btn = QPushButton("📂")
        browse_tpl_btn.setFixedSize(30, 20)
        browse_tpl_btn.clicked.connect(lambda: self.browse_file(self.word_template_path, "Word (*.docx)"))
        
        tpl_layout = QHBoxLayout()
        tpl_layout.addWidget(self.word_template_path)
        tpl_layout.addWidget(browse_tpl_btn)
        
        # 参考代码库 (作弊条) 选择
        self.ref_excel_path = QLineEdit()
        self.ref_excel_path.setPlaceholderText("选择包含参考代码的 Excel (用于填充报告)...")
        browse_ref_btn = QPushButton("📂")
        browse_ref_btn.setFixedSize(30, 20)
        browse_ref_btn.clicked.connect(lambda: self.browse_file(self.ref_excel_path, "Excel (*.xlsx)"))
        
        ref_layout = QHBoxLayout()
        ref_layout.addWidget(self.ref_excel_path)
        ref_layout.addWidget(browse_ref_btn)

        self.gen_report_cb = QCheckBox("生成代码同时生成 Word 报告")
        self.gen_report_cb.setChecked(True)

        report_layout.addRow("Word 模板:", tpl_layout)
        report_layout.addRow("参考代码库:", ref_layout)
        report_layout.addRow("", self.gen_report_cb)
        report_group.setLayout(report_layout)
        layout.addWidget(report_group, 0) # 插入到布局

        # 2. 任务列表区
        task_group = QGroupBox("3. 执行队列")
        tg_layout = QVBoxLayout()

        #全选、反选、全不选按钮
        selection_layout = QHBoxLayout()
        selection_layout.addStretch() # Push buttons to the right
        
        self.btn_sel_all = QPushButton("☑ 全选")
        self.btn_sel_all.setFixedSize(60, 24)
        self.btn_sel_all.setStyleSheet("font-size: 11px; padding: 2px;")
        self.btn_sel_all.clicked.connect(lambda: self.batch_change_selection("all"))
        
        self.btn_sel_inv = QPushButton("🔄 反选")
        self.btn_sel_inv.setFixedSize(60, 24)
        self.btn_sel_inv.setStyleSheet("font-size: 11px; padding: 2px;")
        self.btn_sel_inv.clicked.connect(lambda: self.batch_change_selection("invert"))
        
        self.btn_sel_none = QPushButton("☐ 全不选")
        self.btn_sel_none.setFixedSize(60, 24)
        self.btn_sel_none.setStyleSheet("font-size: 11px; padding: 2px;")
        self.btn_sel_none.clicked.connect(lambda: self.batch_change_selection("none"))
        
        selection_layout.addWidget(self.btn_sel_all)
        selection_layout.addWidget(self.btn_sel_inv)
        selection_layout.addWidget(self.btn_sel_none)
        
        tg_layout.addLayout(selection_layout)
        
        self.task_table = QTableWidget()
        self.task_table.setColumnCount(5) 
        self.task_table.setHorizontalHeaderLabels(["选择", "ID", "需求名称", "输出路径 (相对)", "状态"])
        
        # 设置列宽
        self.task_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Fixed)
        self.task_table.setColumnWidth(0, 50)
        self.task_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.task_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.task_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self.task_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        
        tg_layout.addWidget(self.task_table)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        tg_layout.addWidget(self.progress_bar)
        
        run_layout = QHBoxLayout()
        self.run_btn = QPushButton("▶ 开始流水线生成")
        self.run_btn.setStyleSheet("background-color: #07C160; color: white; font-weight: bold; padding: 10px;")
        self.run_btn.clicked.connect(self.start_pipeline)
        
        self.stop_btn = QPushButton("⏹ 停止")
        self.stop_btn.setEnabled(False)
        self.stop_btn.clicked.connect(self.stop_pipeline)
        
        run_layout.addWidget(self.run_btn)
        run_layout.addWidget(self.stop_btn)
        tg_layout.addLayout(run_layout)
        
        task_group.setLayout(tg_layout)
        layout.addWidget(task_group, 2)

    def batch_change_selection(self, mode):
        """
        Batch update checkbox states in the first column of the task table.
        mode: 'all' (Check all), 'none' (Uncheck all), 'invert' (Toggle state)
        """
        row_count = self.task_table.rowCount()
        for row in range(row_count):
            item = self.task_table.item(row, 0)
            if item: # Ensure the item exists
                if mode == "all":
                    item.setCheckState(Qt.Checked)
                elif mode == "none":
                    item.setCheckState(Qt.Unchecked)
                elif mode == "invert":
                    current_state = item.checkState()
                    new_state = Qt.Unchecked if current_state == Qt.Checked else Qt.Checked
                    item.setCheckState(new_state)

    def select_variable_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择变量表 Excel", "", "Excel Files (*.xlsx *.xls)")
        if path:
            self.var_path_edit.setText(path)
            self.app.config.variable_excel_path = path
            self.app.config.save_config()
            self.var_manager = VariableManager(path)

    def browse_file(self, line_edit, filters):
        path, _ = QFileDialog.getOpenFileName(self, "选择文件", "", filters)
        if path:
            line_edit.setText(path)

    # === 新增：解析 Excel 文件的逻辑 ===
    def parse_excel_file(self, path):
        """解析指定格式的 Excel 需求文件"""
        count = 0
        try:
            df = pd.read_excel(path, engine='openpyxl' if path.endswith('.xlsx') else None)
            # 简单清洗：去除全空行，去除列名空格
            df.dropna(how='all', inplace=True)
            df.columns = [str(col).strip() for col in df.columns]
            
            # 检查必须的列是否存在
            required_cols = ["需求ID", "需求名称"] # 最低要求
            for col in required_cols:
                if col not in df.columns:
                    # 尝试模糊匹配 (不区分大小写)
                    matched = False
                    for existing_col in df.columns:
                        if col.lower() == existing_col.lower():
                            df.rename(columns={existing_col: col}, inplace=True)
                            matched = True
                            break
                    if not matched:
                        print(f"Excel {path} 缺少列: {col}，跳过解析。")
                        return 0

            for index, row in df.iterrows():
                # 提取数据
                req_id = str(row.get("需求ID", "")).strip()
                if not req_id or req_id.lower() == "nan": continue
                
                req_name = str(row.get("需求名称", "")).strip()
                output_files_raw = str(row.get("输出文件", "")).strip()
                req_content = str(row.get("需求内容", "")).strip()
                ref_code = str(row.get("参考代码", "")).strip()
                
                if ref_code.lower() == "nan": ref_code = "无"
                
                # 处理输出文件列表
                target_files = []
                if output_files_raw and output_files_raw.lower() != "nan":
                    # 支持换行、逗号分隔
                    potential_files = re.split(r'[,\s\n]+', output_files_raw)
                    for f in potential_files:
                        f = f.strip()
                        if f:
                            if f.startswith("/") or f.startswith("\\"): f = f[1:]
                            target_files.append(f)
                
                # 兜底文件名
                if not target_files:
                    safe_name = req_id.replace(".", "_").replace(":", "_").strip()
                    target_files = [f"src/auto_generated/{safe_name}.c"]
                
                # 构造数据字典
                parsed_data = {
                    "req_id": req_id,
                    "req_name": req_name,
                    "req_content": req_content,
                    "req_vars": "None", # Excel 模式下暂时假设变量通过 RAG/Var表获取
                    "req_ref_code": ref_code,
                    "output_file": target_files[0],
                    "target_files": target_files
                }
                
                task = BatchTask(f"[Excel Source] {path} Row {index}", parsed_data)
                self.tasks.append(task)
                self.add_task_to_ui(task)
                count += 1
                
        except Exception as e:
            QMessageBox.warning(self, "Excel 解析错误", f"解析 {path} 失败:\n{str(e)}")
            
        return count

    # === 新增：将任务添加到 UI 表格的通用方法 ===
    def add_task_to_ui(self, task):
        row = self.task_table.rowCount()
        self.task_table.insertRow(row)
        
        # CheckBox
        check_item = QTableWidgetItem()
        check_item.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled)
        check_item.setCheckState(Qt.Checked) 
        check_item.setTextAlignment(Qt.AlignCenter) 
        self.task_table.setItem(row, 0, check_item)
        
        # Data
        self.task_table.setItem(row, 1, QTableWidgetItem(task.id))
        self.task_table.setItem(row, 2, QTableWidgetItem(task.name))
        
        files_display = "\n".join(task.target_files)
        self.task_table.setItem(row, 3, QTableWidgetItem(files_display))
        self.task_table.setItem(row, 4, QTableWidgetItem(task.status))
        
        self.task_table.resizeRowToContents(row)

    def parse_tasks(self):
        raw_input = self.batch_input_edit.toPlainText().strip()
        if not raw_input:
            return
            
        full_text = ""
        # 1. 清空旧列表
        self.tasks = []
        self.task_table.setRowCount(0)
        
        lines = raw_input.split('\n')
        excel_task_count = 0
        text_files_count = 0
        
        # === 遍历输入行，区分 Excel、文本文件、纯文本 ===
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith("file:///"):
                line = unquote(line[8:])
            
            path = line.strip('"').strip("'")
            if re.match(r'^/[a-zA-Z]:', path):
                path = path[1:]
                
            if os.path.isfile(path):
                # 判定文件类型
                if path.lower().endswith(('.xlsx', '.xls')):
                    # === 分支 A: Excel 解析 ===
                    cnt = self.parse_excel_file(path)
                    excel_task_count += cnt
                else:
                    # === 分支 B: 文本/MD 读取 ===
                    try:
                        with open(path, 'r', encoding='utf-8') as f:                      
                            full_text += f.read() + "\n\n" 
                            text_files_count += 1
                    except Exception as e:
                        print(f"读取失败: {path}")
            else:
                # === 分支 C: 纯文本输入 ===
                # 只有当没有识别出任何文件读取时，才把这一行当作需求文本的一部分
                # 或者，如果这行看起来像路径但文件不存在，也拼进去防止误判
                full_text += line + "\n"

        self.app.status_label.setText(f"解析结果: Excel任务 {excel_task_count} 个, 文本文件 {text_files_count} 个")
        
        # === 处理合并后的文本内容 (正则解析) ===
        regex_task_count = 0
        if full_text.strip():
            pattern = r'(需求id[：:].*?)(?=\n\s*需求id[：:]|\Z)'
            matches = re.finditer(pattern, full_text, re.IGNORECASE | re.DOTALL)
            
            for match in matches:
                block = match.group(1).strip()
                if not block: continue
                
                data = self.app.parse_requirement_text(block)
                
                target_files = []
                # 正则捕获内容 (跨行抓取)
                file_pattern = r"输出文件[：:]\s*(.*?)(?=\n\s*(?:需求内容|变量|参考代码|阶段|类型|需求id)|$)"
                match_out = re.search(file_pattern, block, re.IGNORECASE | re.DOTALL)
                
                if match_out:
                    raw_content = match_out.group(1).strip()
                    potential_files = re.split(r'[,\s\n]+', raw_content)
                    for f in potential_files:
                        f = f.strip()
                        if f: 
                            if f.startswith("/") or f.startswith("\\"):
                                f = f[1:]
                            target_files.append(f)

                if not target_files:
                    safe_name = data["req_id"].replace(".", "_").replace(":", "_").strip()
                    target_files = [f"src/auto_generated/{safe_name}.c"]

                # 补全 parsed_data 以复用 logic
                data["target_files"] = target_files
                data["output_file"] = target_files[0] if target_files else ""

                task = BatchTask(block, data)
                
                if data.get("req_id", "N/A") != "N/A":
                    self.tasks.append(task)
                    self.add_task_to_ui(task)
                    regex_task_count += 1
            
        total = excel_task_count + regex_task_count
        if total > 0:
            QMessageBox.information(self, "解析完成", f"共识别到 {total} 个任务。\n- Excel来源: {excel_task_count}\n- 文本解析: {regex_task_count}")
        else:
            QMessageBox.warning(self, "解析失败", "未识别到有效任务。请检查文件格式或文本内容。")

    def clear_tasks(self):
        self.tasks = []
        self.task_table.setRowCount(0)
        self.progress_bar.setValue(0)

    def start_pipeline(self):
        if not self.tasks:
            return
            
        root = self.app.config.project_root
        if not root:
            QMessageBox.warning(self, "未配置根目录", "请先在【系统设置】中配置工程根目录！")
            return
        
        excel_path = self.var_path_edit.text().strip()
        if excel_path and os.path.exists(excel_path):
            if not hasattr(self, 'var_manager') or getattr(self, 'current_excel_path', '') != excel_path:
                self.var_manager = VariableManager(excel_path)
                self.current_excel_path = excel_path
        else:
            self.var_manager = None

        # === 新增: 初始化参考代码管理器 ===
        ref_excel = self.ref_excel_path.text().strip()
        if ref_excel and os.path.exists(ref_excel):
            self.ref_manager = RefCodeManager(ref_excel)
        else:
            self.ref_manager = None

        # === 新增: 准备临时文件夹 ===
        self.generated_temp_files = [] 
        self.temp_dir = os.path.join(self.app.config.project_root, "temp_reports")
        os.makedirs(self.temp_dir, exist_ok=True)
            
        self.is_running = True
        self.run_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.current_task_index = -1
        self.progress_bar.setMaximum(len(self.tasks))
        self.progress_bar.setValue(0)
        
        self.run_next_task()

    def stop_pipeline(self):
        self.is_running = False
        
        if hasattr(self, 'worker') and self.worker.isRunning():
            try:
                self.worker.finished_signal.disconnect()
                self.worker.error_signal.disconnect()
            except Exception:
                pass
            self.worker.terminate()
            self.worker.wait(100)

        if 0 <= self.current_task_index < len(self.tasks):
            self.update_task_status(self.current_task_index, "⚠️ 已停止 (用户终止)")
            item = self.task_table.item(self.current_task_index, 4)
            if item:
                item.setForeground(Qt.red)

        self.run_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.app.status_label.setText("流水线已手动停止")
        QMessageBox.information(self, "提示", "流水线已强制停止。")

    def run_next_task(self):
        if not self.is_running:
            return
            
        self.current_task_index += 1
        
        if self.current_task_index >= len(self.tasks):
            self.finish_pipeline()
            return
            
        task = self.tasks[self.current_task_index]
        
        # CheckBox 跳过逻辑
        check_item = self.task_table.item(self.current_task_index, 0)
        if check_item and check_item.checkState() == Qt.Unchecked:
            self.update_task_status(self.current_task_index, "⏭️ 已跳过")
            QTimer.singleShot(10, self.run_next_task)
            return
        
        self.update_task_status(self.current_task_index, "🔍 检索上下文(RAG/Vars)...") 
        self.task_table.scrollToItem(self.task_table.item(self.current_task_index, 0))
        QApplication.processEvents() 
        
        # =========================================================
        # 第一阶段：RAG 向量知识库检索
        # =========================================================
        rag_context = ""
        context_parts = []
        
        if self.pipe_use_rag_cb.isChecked():
            selected_kbs = []
            count = self.pipe_kb_list.count()
            for i in range(count):
                item = self.pipe_kb_list.item(i)
                if item.checkState() == Qt.Checked:
                    selected_kbs.append(item.text())
            
            if selected_kbs:
                full_query = f"{task.id} {task.name}\n{task.content}"
                limit_length = 800 
                query_text = full_query[:limit_length] if len(full_query) > limit_length else full_query
                
                for kb_name in selected_kbs:
                    try:
                        res = self.app.rag_manager.recall(query_text, kb_name)
                        if res and "未找到" not in res and "检索出错" not in res:
                             context_parts.append(f"【参考规范/文档来源: {kb_name}】\n{res}")
                    except Exception as e:
                        print(f"KB {kb_name} 检索失败: {e}")
            
                if context_parts:
                    rag_context = "\n\n".join(context_parts)
                else:
                    rag_context = "（已启用 RAG，但未在选定知识库中检索到高相关性内容）"

        # =========================================================
        # 第二阶段：Excel 变量表匹配
        # =========================================================
        matched_vars_text = "无 (未配置变量表或无匹配项)"
        
        if hasattr(self, 'var_manager') and self.var_manager and self.var_manager.is_loaded:
            search_query = f"{task.name} {task.content}"
            relevant_vars = self.var_manager.search_relevant_vars(search_query, top_k=15)
            
            if relevant_vars:
                matched_vars_text = "\n".join(relevant_vars)
                task.used_global_vars = matched_vars_text

        variable_instruction = (
            "\n【变量使用与推理指南 (Variable Logic)】:\n"
            "系统已根据需求检索到了以下可能相关的【现有变量库】：\n"
            "\"\"\"\n"
            f"{matched_vars_text}\n"
            "\"\"\"\n"
            "**你的行动指令**：\n"
            "1. **优先匹配**：仔细检查上述列表。如果列表中存在符合需求的变量（名称或物理含义匹配），**必须直接使用**该变量的 ID 和 数据类型，严禁重复定义。\n"
            "2. **推理生成**：如果上述列表中没有你需要的变量，请**参考列表中的命名风格**（如前缀 ELEC/HYD、下划线用法、大写习惯），自行推理并生成新的变量 ID 和定义。\n"
            "3. **一致性**：确保新生成的变量与现有变量在风格上保持高度一致。\n"
        )

        # =========================================================
        # 第三阶段：构建 Prompt 规则
        # =========================================================
        external_rules = ""
        rule_path = self.app.rule_path_edit.text().strip()
        if rule_path and os.path.exists(rule_path):
            try:
                with open(rule_path, 'r', encoding='utf-8') as f:
                    external_rules = f"【用户指定规则文件】:\n{f.read()}\n"
            except Exception:
                pass

        seven_step_pipeline = (
            "\n【代码生成-标准作业流程 (SOP)】:\n"
            "在编写代码实现该需求时，必须严格在函数内部体现以下 7 个阶段的逻辑（可使用注释标记阶段）：\n"
            "1. **开始 (Start)**: 变量定义、静态变量初始化、结构体清零。\n"
            "2. **信号处理 (Signal Proc)**: 对输入数据（传感器/参数）进行滤波、有效性校验、归一化。\n"
            "3. **能源管理 (Power Mgmt)**: 检查当前电源状态、功耗模式、是否允许执行高功耗操作。\n"
            "4. **系统控制 (Control Algo)**: 执行核心业务逻辑（如 PID 计算、状态机流转、阈值判定）。\n"
            "5. **告警输出 (Alarm Out)**: 检查上述逻辑是否触发故障，置位错误码或故障标志。\n"
            "6. **信号输出 (Signal Out)**: 将最终结果写入硬件寄存器、输出变量或缓冲区。\n"
            "7. **结束 (End)**: 更新历史状态变量（如上一次误差）、返回执行结果。\n"
        )

        files_instruction = ""
        if task.target_files and len(task.target_files) > 0:
            files_instruction = "\n【目标文件列表】(请依次生成):\n"
            for f in task.target_files:
                files_instruction += f"- {f}\n"
            
            files_instruction += (
                "\n【重要：多文件分割格式】\n"
                "请严格按照以下格式输出每个文件的代码，不要合并：\n"
                "/* === FILE_START: src/xxx/filename.c === */\n"
                "...(代码内容)...\n"
                "/* === FILE_END === */\n\n"
            )

        pipeline_constraint = (
            "\n【流水线模式-强制格式约束】:\n"
            "1. 请直接输出函数实现代码，严禁包含 main() 函数。\n"
            "2. 代码必须是完整的、可编译的 C 函数。\n"
            "3. **必须在代码中用注释标出上述 7 个阶段** (例如: /* 1. 开始 */ ...)。\n"
        )

        # 处理参考代码
        ref_code_section = ""
        if hasattr(task, 'ref_code') and task.ref_code and task.ref_code != "无" and task.ref_code != "None":
            ref_code_section = f"\n【参考代码/遗留代码】:\n{task.ref_code}\n"
        else:
            ref_code_section = "无 (新开发需求)"

        base_rule_header = "遵循 GJB 5369-2005 通用安全标准。\n" if not external_rules else external_rules
        final_rules = f"{base_rule_header}\n{variable_instruction}\n{seven_step_pipeline}\n{files_instruction}\n{pipeline_constraint}"

        # 构造 req_vars
        parsed_local_vars = getattr(task, 'vars', "").strip()
        if parsed_local_vars and parsed_local_vars != "None" and parsed_local_vars != "无":
            final_req_vars = f"{parsed_local_vars}\n(注：请同时结合 Rules 中检索到的系统变量进行开发)"
        else:
            if matched_vars_text and "未配置变量表" not in matched_vars_text:
                final_req_vars = "本需求未定义局部变量，**请严格使用 [Rules] - [变量使用与推理指南] 章节中列出的系统变量**。"
            else:
                final_req_vars = "无明确变量定义，请根据逻辑自行推断（需符合命名规范）。"

        # =========================================================
        # 第四阶段：调用 LLM
        # =========================================================
        try:
            prompt = self.app.config.prompt_template.format(
                req_id=task.id,
                req_name=task.name,
                req_content=task.content,
                req_vars=final_req_vars,
                req_ref_code=ref_code_section, 
                context=rag_context,
                rules=final_rules
            )
        except KeyError:
            prompt = (
                f"# Task: {task.name}\n# Requirement:\n{task.content}\n"
                f"# Rules:\n{final_rules}\n# Context:\n{rag_context}\n"
                f"Please write code."
            )

        self.update_task_status(self.current_task_index, "生成中(AI思考)...")
        
        self.worker = GenerationThread(
            self.app.config.api_url,
            self.app.config.api_key,
            self.app.config.model_name,
            prompt
        )
        self.worker.finished_signal.connect(self.on_task_finished)
        self.worker.error_signal.connect(self.on_task_error)
        self.worker.start()

    def on_task_finished(self, text):
        if not self.is_running: return
        
        task = self.tasks[self.current_task_index]
        root = self.app.config.project_root
        final_code_display = "" 
        
        import re
        # 定义你的多文件分割正则
        pattern = r"/\* === FILE_START: (.*?) === \*/(.*?)/\* === FILE_END === \*/"
        matches = re.findall(pattern, text, re.DOTALL)
        
        success_count = 0
        status_msg = ""
        
        if matches:
            # === 情况 A：匹配到了多文件格式 ===
            # 1. 遍历写入文件
            for filename, content in matches:
                filename = filename.strip()
                content = content.strip()
                
                # 拼接代码用于 Word 报告展示
                if filename.lower().endswith(".c"):
                    final_code_display += f"{content}\n\n"
                
                try:
                    full_path = os.path.join(root, filename)
                    os.makedirs(os.path.dirname(full_path), exist_ok=True)
                    with open(full_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    success_count += 1
                except Exception as e:
                    print(f"写入 {filename} 失败: {e}")
            
            status_msg = f"完成 (生成 {success_count} 个文件)"
            
            # ⚠️ 关键修正：将正则提取到的内容赋值给任务对象，用于 Word 报告
            task.generated_clean_code = final_code_display

        else:
            # === 情况 B：没匹配到格式，尝试回退到普通 Markdown 提取 ===
            clean_code = self.app.extract_code_blocks(text)
            task.generated_clean_code = clean_code # 使用通用提取结果
            
            # 尝试作为单文件写入
            target = task.output_rel_path
            if not target and task.target_files:
                target = task.target_files[0]
            
            if target and "未检测到" not in clean_code: # 只有提取成功才写入
                try:
                    full_path = os.path.join(root, target)
                    os.makedirs(os.path.dirname(full_path), exist_ok=True)
                    with open(full_path, 'w', encoding='utf-8') as f:
                        f.write(clean_code)
                    status_msg = "完成 (单文件)"
                except Exception as e:
                    status_msg = f"写入失败: {e}"
            else:
                status_msg = "完成 (格式不匹配)"

        # === 新增: 生成临时 Word 报告 ===
        if self.gen_report_cb.isChecked():
            template_path = self.word_template_path.text().strip()
            if template_path and os.path.exists(template_path):
                # 构造临时文件名 (带序号)
                safe_name = task.id.replace(".", "_").strip()
                temp_doc_name = f"{self.current_task_index:03d}_{safe_name}.docx"
                temp_path = os.path.join(self.temp_dir, temp_doc_name)
                
                # 获取参考代码
                ref_code_str = "无"
                if hasattr(self, 'ref_manager') and self.ref_manager:
                    ref_code_str = self.ref_manager.get_code(task.id) # 从GT表拿
                elif task.ref_code and task.ref_code != "None":
                    ref_code_str = task.ref_code # 备用

                raw_vars = getattr(task, 'used_global_vars', "")
                final_vars_str = raw_vars if raw_vars and raw_vars.strip() else "无"

                # 填入数据
                data = {
                    "{需求名称}": task.name,
                    "{需求ID}": task.id,
                    "{需求内容}": task.content,
                    "{参考代码}": ref_code_str,
                    "{实际输出代码}": task.generated_clean_code,
                    "{检索的全局变量}": final_vars_str
                }
                
                success, path = WordReportGenerator.generate_report(template_path, temp_path, data)
                if success:
                    self.generated_temp_files.append(path)
            
        self.update_task_status(self.current_task_index, status_msg)
        self.progress_bar.setValue(self.current_task_index + 1)
        self.run_next_task()

    def on_task_error(self, err_msg):
        if not self.is_running: return
        self.update_task_status(self.current_task_index, f"API错误: {err_msg}")
        self.run_next_task()

    def update_task_status(self, row, status):
        self.task_table.setItem(row, 4, QTableWidgetItem(status))
        self.task_table.scrollToItem(self.task_table.item(row, 0))

    def finish_pipeline(self):
        self.is_running = False
        self.run_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        if self.gen_report_cb.isChecked() and self.generated_temp_files:
            self.app.status_label.setText("正在合并/追加 Word 报告...")
            QApplication.processEvents()
            
            # 设置固定的总报告文件名
            final_report_name = "Project_Total_Report.docx" 
            final_path = os.path.join(self.app.config.project_root, "reports", final_report_name)
            os.makedirs(os.path.dirname(final_path), exist_ok=True)
            
            # 调用合并 (会自动判断是追加还是新建)
            success, msg = WordReportGenerator.merge_reports(self.generated_temp_files, final_path)
            
            if success:
                # 删除临时目录
                try:
                    shutil.rmtree(self.temp_dir)
                except:
                    pass
                QMessageBox.information(self, "完成", f"全部完成！报告已更新至:\n{final_path}")
            else:
                QMessageBox.warning(self, "警告", f"合并出错: {msg}")
        else:
            QMessageBox.information(self, "完成", "流水线队列处理完毕！")
